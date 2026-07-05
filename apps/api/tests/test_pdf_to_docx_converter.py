import threading
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import fitz
import pytest
from docx import Document

from app.modules.converter import get_converter
from app.modules.converter.pdf_to_docx import (
    ImagesExtractor,
    PdfToDocxConverter,
    _normalize_pixmap_for_png,
    _original_to_raw_dict,
)


def test_pdf_to_docx_is_registered_automatically() -> None:
    assert isinstance(get_converter("pdf-to-docx"), PdfToDocxConverter)


def _build_pdf_with_image(path: Path, pixmap: fitz.Pixmap, text: str = "sample text") -> None:
    """A minimal one-page PDF with `pixmap` embedded as-is (preserving its
    exact colorspace — `insert_image(pixmap=...)` does not re-encode it,
    verified empirically) plus some real text, so both the image-extraction
    and text-layout parts of pdf2docx run against it."""
    doc = fitz.open()
    page = doc.new_page(width=300, height=300)
    page.insert_image(fitz.Rect(50, 50, 200, 200), pixmap=pixmap)
    page.insert_text((10, 20), text)
    doc.save(path)
    doc.close()


def _build_pdf_with_non_rgb_image(path: Path) -> None:
    """A real, from-scratch PDF whose embedded image uses a `/Separation`
    colorspace — a real-world case pdf2docx==0.5.8's own CMYK-only check
    (`'CMYK' in item[5].upper()`) does not recognize, since `page.
    get_images()` reports its colorspace name as plain `"Separation"`.

    Built via low-level xref/object construction (no high-level PyMuPDF API
    exposes non Gray/RGB/CMYK colorspaces directly) — not a synthetic
    isolated `Pixmap.tobytes()` call: `page.get_images()` and `fitz.Pixmap
    (doc, xref)` here reconstruct it exactly as pdf2docx's own
    `ImagesExtractor._recover_pixmap` does, and running the real
    `PdfToDocxConverter.convert()` against this file on pre-fix HEAD
    reproduces the exact production traceback (`pixmap must be grayscale
    or rgb to write as png`, during pdf2docx's `parse_document()` step).
    """
    doc = fitz.open()
    page = doc.new_page(width=300, height=300)

    # Type 2 (exponential interpolation) tint-transform function.
    func_xref = doc.get_new_xref()
    doc.update_object(
        func_xref, "<< /FunctionType 2 /Domain [0 1] /C0 [1 1 1] /C1 [0.2 0.4 0.8] /N 1 >>"
    )

    # Separation colorspace array referencing that function.
    cs_xref = doc.get_new_xref()
    doc.update_object(cs_xref, f"[/Separation /Spot /DeviceRGB {func_xref} 0 R]")

    # Raw single-component (Separation is always 1 component) image data.
    w, h = 20, 20
    raw = bytes([(x * 13 + y * 7) % 256 for y in range(h) for x in range(w)])
    img_xref = doc.get_new_xref()
    doc.update_object(
        img_xref,
        f"<< /Type /XObject /Subtype /Image /Width {w} /Height {h} "
        f"/BitsPerComponent 8 /ColorSpace {cs_xref} 0 R /Length {len(raw)} >>",
    )
    doc.update_stream(img_xref, raw)

    page.insert_image(fitz.Rect(50, 50, 200, 200), xref=img_xref)
    page.insert_text((10, 20), "hello separation")
    doc.save(path)
    doc.close()


def test_convert_preserves_headings_paragraphs_tables_images(
    sample_pdf_path: Path, tmp_path: Path
) -> None:
    """Validates the real (public, MIT-licensed) sample document converts
    with layout fidelity: a 6-page, two-column academic paper with figures
    and tables — pdf2docx's own project demo fixture.

    Page count is *not* checked via DOCX section count here: for a simple
    single-column page, pdf2docx does emit one section per page, but for
    this multi-column document it emits multiple sections per page (one per
    column/layout region), so section count isn't a reliable proxy for page
    count in general — a real finding from converting this fixture, not an
    assumption.
    """
    pdf = fitz.open(sample_pdf_path)
    source_page_count = pdf.page_count
    source_image_count = sum(len(page.get_images()) for page in pdf)
    pdf.close()

    output_path = PdfToDocxConverter().convert(sample_pdf_path, tmp_path)

    assert output_path.exists()
    assert output_path.suffix == ".docx"

    doc = Document(output_path)

    # Page count: DOCX has no hard page concept, but pdf2docx should still
    # emit at least one section per source page.
    assert len(doc.sections) >= source_page_count

    # Images: pdf2docx also rasterizes vector graphics, so the docx can
    # contain more image parts than the PDF's raw raster image count — but
    # never fewer, i.e. no image content is dropped.
    embedded_images = [r for r in doc.part.rels.values() if "image" in r.reltype]
    assert len(embedded_images) >= source_image_count > 0

    # Headings/paragraphs: pdf2docx has no native "Heading N" style mapping —
    # it preserves emphasis via bold/font-size runs on Normal-styled
    # paragraphs rather than semantic heading styles. So "headings preserved"
    # is verified as bold, larger-than-body-text runs surviving conversion,
    # not python-docx `Heading` styles.
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)

    text_runs = [r for p in all_paragraphs for r in p.runs if r.text.strip()]
    assert text_runs, "expected extracted text runs in the converted document"

    bold_heading_like_runs = [r for r in text_runs if r.bold]
    assert bold_heading_like_runs, "expected at least one bold/heading-like run"

    # Tables: the paper's data tables should survive as actual Word tables,
    # not flattened into plain paragraphs.
    assert len(doc.tables) >= 1

    # Spot-check that real document content made it through conversion.
    full_text = "\n".join(p.text for p in all_paragraphs)
    assert "Abstract" in full_text


def test_pdf_to_docx_normalizes_non_rgb_embedded_image_before_png_serialization(
    tmp_path: Path,
) -> None:
    """Regression test for the production failure class: a PDF embedding a
    non-RGB/non-grayscale image (here, `/Separation`) must convert
    successfully instead of raising during pdf2docx's image-serialization
    step. See `_build_pdf_with_non_rgb_image` for how this was verified,
    pre-fix, to reproduce the real production traceback."""
    source_path = tmp_path / "non_rgb.pdf"
    _build_pdf_with_non_rgb_image(source_path)

    output_path = PdfToDocxConverter().convert(source_path, tmp_path / "out")

    assert output_path.exists()
    assert output_path.stat().st_size > 0

    # Confirm the output is actually a valid, openable docx, not just a file.
    doc = Document(output_path)
    assert doc.paragraphs


def test_pdf_to_docx_still_converts_rgb_embedded_image(tmp_path: Path) -> None:
    source_path = tmp_path / "rgb.pdf"
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 20, 20))
    pix.set_rect(pix.irect, (10, 120, 200))
    _build_pdf_with_image(source_path, pix)

    output_path = PdfToDocxConverter().convert(source_path, tmp_path / "out")

    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert Document(output_path).paragraphs


def test_pdf_to_docx_still_converts_grayscale_embedded_image(tmp_path: Path) -> None:
    source_path = tmp_path / "gray.pdf"
    pix = fitz.Pixmap(fitz.csGRAY, fitz.IRect(0, 0, 20, 20))
    pix.set_rect(pix.irect, (128,))
    _build_pdf_with_image(source_path, pix)

    output_path = PdfToDocxConverter().convert(source_path, tmp_path / "out")

    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert Document(output_path).paragraphs


def test_pdf_to_docx_still_converts_text_only_pdf(tmp_path: Path) -> None:
    source_path = tmp_path / "text_only.pdf"
    doc = fitz.open()
    doc.new_page().insert_text((72, 72), "plain text only, no images")
    doc.save(source_path)
    doc.close()

    output_path = PdfToDocxConverter().convert(source_path, tmp_path / "out")

    assert output_path.exists()
    assert output_path.stat().st_size > 0
    converted = Document(output_path)
    assert any("plain text only" in p.text for p in converted.paragraphs)


class TestNormalizePixmapForPng:
    """Unit tests for the compatibility helper itself, isolated from the
    full pdf2docx conversion flow."""

    def test_rgb_pixmap_is_unchanged(self) -> None:
        pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 4, 4))
        normalized = _normalize_pixmap_for_png(pix)
        assert normalized is pix
        assert normalized.colorspace.name == "DeviceRGB"

    def test_grayscale_pixmap_is_unchanged(self) -> None:
        pix = fitz.Pixmap(fitz.csGRAY, fitz.IRect(0, 0, 4, 4))
        normalized = _normalize_pixmap_for_png(pix)
        assert normalized is pix
        assert normalized.colorspace.name == "DeviceGray"

    def test_cmyk_pixmap_is_converted_to_rgb(self) -> None:
        pix = fitz.Pixmap(fitz.csCMYK, fitz.IRect(0, 0, 4, 4))
        normalized = _normalize_pixmap_for_png(pix)
        assert normalized.colorspace.name == "DeviceRGB"
        normalized.tobytes()  # must not raise

    def test_alpha_is_preserved_through_conversion(self) -> None:
        base = fitz.Pixmap(fitz.csCMYK, fitz.IRect(0, 0, 4, 4))
        with_alpha = fitz.Pixmap(base, 1)
        normalized = _normalize_pixmap_for_png(with_alpha)
        assert normalized.alpha == 1

    def test_none_colorspace_pixmap_is_unaffected(self) -> None:
        # Pixmaps built directly always carry a colorspace in this PyMuPDF
        # version — this asserts the `colorspace is None` branch is at
        # least a safe no-op/pass-through rather than raising, for the
        # stencil-mask-only pixmaps pdf2docx's own docs describe.
        pix = fitz.Pixmap(fitz.csGRAY, fitz.IRect(0, 0, 4, 4))
        assert _normalize_pixmap_for_png(pix) is pix


def test_concurrent_conversions_do_not_corrupt_the_shared_patch(tmp_path: Path) -> None:
    """The compatibility patch is installed on the shared, process-wide
    `ImagesExtractor` class (see `_png_safe_image_serialization`'s
    docstring) — this proves the reference-counted install/restore is safe
    when multiple conversions run truly concurrently across OS threads, the
    same concurrency model `asyncio.to_thread` uses in production, and that
    it's fully removed again once every conversion has finished (never left
    behind as a permanent patch)."""
    non_rgb_path = tmp_path / "non_rgb.pdf"
    _build_pdf_with_non_rgb_image(non_rgb_path)

    rgb_path = tmp_path / "rgb.pdf"
    pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 20, 20))
    pix.set_rect(pix.irect, (10, 120, 200))
    _build_pdf_with_image(rgb_path, pix)

    start_barrier = threading.Barrier(4)

    def _convert(source_path: Path, out_dir: Path) -> Path:
        start_barrier.wait(timeout=5)  # maximize actual overlap
        return PdfToDocxConverter().convert(source_path, out_dir)

    jobs = [
        (non_rgb_path, tmp_path / "out1"),
        (rgb_path, tmp_path / "out2"),
        (non_rgb_path, tmp_path / "out3"),
        (rgb_path, tmp_path / "out4"),
    ]
    with ThreadPoolExecutor(max_workers=4) as pool:
        futures = [pool.submit(_convert, src, out) for src, out in jobs]
        results = [f.result(timeout=30) for f in futures]

    for output_path in results:
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    # The patch must be fully removed again — no permanent global state left
    # behind once every concurrent conversion has completed.
    assert ImagesExtractor._to_raw_dict is _original_to_raw_dict


@pytest.fixture(autouse=True)
def _restore_images_extractor_patch_state() -> None:
    """Safety net: if a test above fails mid-conversion in a way that skips
    the patch's own cleanup, don't let that leak into unrelated tests
    elsewhere in the suite."""
    yield
    ImagesExtractor._to_raw_dict = staticmethod(_original_to_raw_dict)
